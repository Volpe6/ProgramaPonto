import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

class Ponto(object):
    '''
        Adiciona ponto em arquivo do word
    '''

    DIR_PADRAO = f'{Path().cwd()}/pontos'

    MESES = (
        'janeiro', 
        'fevereiro', 
        'marco', 
        'abril', 
        'maio', 
        'junho', 
        'julho', 
        'agosto', 
        'setembro', 
        'outubro',
        'novembro',
        'dezembro')

    def __init__(self, nome_bolsista, modalidade_bolsa, observacoes):
        self.nome_bolsista        = nome_bolsista
        self.modalidade_bolsa     = modalidade_bolsa
        self.observacoes          = observacoes
        self.doc                  = None
        self.table                = None
        self.nome_pasta_mes_atual = ''
    
    def processa(self):
        self.gerencia_arq()
        self.abri_arq_word()
        self.conf_padrao()
        self.add_nome_bolsista()
        self.add_nome_modalidade_bolsa()
        self.add_mes_ano()
        self.add_observacao()

        
    def abri_arq_word(self):
        caminho_arq = f'{self.nome_pasta_mes_atual}/ponto-{self.get_mes_atual()}.docx'
        if self.existe_pasta(caminho_arq):
            self.doc   = Document(caminho_arq)
            self.table = self.doc.tables
            return
        self.doc = Document(f'{Path().cwd()}/modelo-ponto.docx')
        self.table = self.doc.tables


    def salva_arq(self):
        self.doc.save(f'{self.nome_pasta_mes_atual}/ponto-{self.get_mes_atual()}.docx')


    def add_hora_entrada(self):
        cell  = self.table[0].cell(self.get_data().day + 2, 1)
        self.add_conteudo_celula(cell, self.get_hora_minuto_atual())


    def add_hora_saida(self):
        cell  = self.table[0].cell(self.get_data().day + 2, 3)
        self.add_conteudo_celula(cell, self.get_hora_minuto_atual())


    def add_nome_bolsista(self):
        cell = self.table[0].cell(0, 0)
        self.add_conteudo_celula_negrito(cell, 'Nome do (a) bolsista: ', self.nome_bolsista)


    def add_nome_modalidade_bolsa(self):
        cell = self.table[0].cell(1, 0)
        self.add_conteudo_celula_negrito(cell, 'Modalidade da bolsa: ', self.modalidade_bolsa)


    def add_mes_ano(self):
        cell    = self.table[0].cell(1, 3)
        mes_ano = f'{self.get_data().month}/{self.get_data().year}' 
        self.add_conteudo_celula_negrito(cell, 'Mês/Ano: ', mes_ano)


    def add_observacao(self):
        cell = self.table[0].cell(34, 1)
        txt  = ''
        for texto in self.observacoes:
            txt += f'{texto}\n'
        self.add_conteudo_celula_negrito(cell, 'Observações: ', txt)


    def add_conteudo_celula(self, cell, conteudo):
        paragrafo           = cell.paragraphs[0]
        paragrafo.text      = ''
        paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragrafo.add_run(conteudo).bold = True


    def add_conteudo_celula_negrito(self, cell, label, conteudo):
        for paragrafo in cell.paragraphs:
            paragrafo.text = ''

        paragrafo      = cell.paragraphs[0]
        paragrafo.text = ''
        paragrafo.add_run(label).bold = True
        paragrafo.add_run(conteudo)


    def conf_padrao(self):
        font      = self.doc.styles['Normal'].font
        font.name = 'Arial'
        font.size = 140000

    def gerencia_arq(self):
        dir_pontos = '{}/pontos'.format(Path().cwd())

        nome_pasta = 'ponto-{}-{}'.format(self.get_mes_atual(), self.get_data().year)

        #verifica se o arquivo de pontos existe
        if not self.existe_pasta(self.DIR_PADRAO):
            self.cria_pasta(self.DIR_PADRAO)
        
        #verifica se existe pasta do mes atual
        if not self.existe_pasta(f'{self.DIR_PADRAO}/{nome_pasta}'):
            self.cria_pasta(f'{self.DIR_PADRAO}/{nome_pasta}')


        self.nome_pasta_mes_atual = f'{dir_pontos}/{nome_pasta}'
            

    def existe_pasta(self, caminho):
        return Path(caminho).exists()


    def cria_pasta(self, caminho):
        Path(caminho).mkdir()
        

    def get_data(self):
        return datetime.datetime.now()


    def get_mes_atual(self):
        data = self.get_data()
        return self.MESES[data.month - 1]

    
    def get_hora_minuto_atual(self):
        data = self.get_data()
        return '{:02d}:{:02d}'.format(data.hour, data.minute)


# doc = Document('C:/Users/Drew/Documents/python/ponto/pontos/ponto-junho-2020/ponto-junho.docx')
# doc = Document('modelo-ponto.docx')

# table = doc.tables

# for td in table:
#     for row in range(len(td.rows)):
#         for col in range(len(td.columns)):
#             print(f'linha: {row}, col: {col}. conteudo: {td.cell(row, col).text}')